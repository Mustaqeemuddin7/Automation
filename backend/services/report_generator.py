from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
import concurrent.futures
from docx.enum.text import WD_BREAK
import pandas as pd

import os


def generate_hod_remark(attendance_percent, cie_percent, backlog_count):
    """Generate HOD remark based on attendance %, CIE marks %, and backlog count.
    
    Rules (checked in order — negative conditions first, then positive):
      < 75% attendance         → Poor Attendance
      < 50% CIE                → Academically Weak
      backlogs >= 5 (only)     → Backlog Concern
      >= 90% att, >= 85% CIE, 0 backlogs   → Outstanding
      >= 85% att, >= 75% CIE, <= 2 backlogs → Very Good
      >= 80% att, >= 70% CIE, <= 2 backlogs → Good Performance
      >= 75% att, >= 60% CIE, 3-4 backlogs  → Satisfactory
      >= 75% att, >= 70% CIE, >= 5 backlogs → Needs Improvement
      else                                   → Satisfactory
    """
    # Negative / warning conditions first (override positives)
    if attendance_percent < 50:
        return 'Poor Attendance'
    if cie_percent < 50:
        return 'Academically Weak'
    
    # Positive conditions (most specific first)
    if attendance_percent >= 90 and cie_percent >= 85 and backlog_count == 0:
        return 'Outstanding'
    if attendance_percent >= 85 and cie_percent >= 75 and backlog_count <= 2:
        return 'Very Good'
    if attendance_percent >= 80 and cie_percent >= 70 and backlog_count <= 2:
        return 'Good Performance'
    if attendance_percent >= 75 and cie_percent >= 60 and 3 <= backlog_count <= 4:
        return 'Satisfactory'
    if attendance_percent >= 75 and cie_percent >= 70 and backlog_count >= 5:
        return 'Needs Improvement'
    if backlog_count >= 5:
        return 'Backlog Concern'
    
    return 'Satisfactory'

def add_logo_and_header(doc, department_name):
    """Add institutional header with logo on left, text on right (table layout), matching main format.docx"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # Create a table for logo (left) + header text (right)
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    
    # Remove table borders except bottom (add bottom border only)
    tbl = header_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        tblBorders.append(border)
    # Add visible bottom border
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')  # Border width
    bottom_border.set(qn('w:space'), '0')
    bottom_border.set(qn('w:color'), '000000')  # Black color
    tblBorders.append(bottom_border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    
    # Logo cell (left) - set cell width explicitly
    logo_cell = header_table.cell(0, 0)
    logo_cell.width = Inches(1.0)
    logo_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'assets', 'image.png')
    if os.path.exists(logo_path):
        try:
            logo_para = logo_cell.paragraphs[0]
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_para.add_run()
            logo_run.add_picture(logo_path, width=Inches(0.8))
        except Exception:
            pass
    
    # Header text cell (right) - set cell width explicitly
    text_cell = header_table.cell(0, 1)
    text_cell.width = Inches(6.5)
    from docx.shared import Twips
    
    # Line 1: Institution name
    p1 = text_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Twips(0)
    p1.paragraph_format.space_after = Twips(0)
    run = p1.add_run("LORDS INSTITUTE OF ENGINEERING &TECHNOLOGY")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    
    # Line 2: Autonomous
    p2 = text_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Twips(0)
    p2.paragraph_format.space_after = Twips(0)
    run = p2.add_run("(Autonomous)")
    run.font.name = 'Aptos Display (Headings)'
    run.font.size = Pt(14)
    
    # Line 3: AICTE
    p3 = text_cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Twips(0)
    p3.paragraph_format.space_after = Twips(0)
    run = p3.add_run("Approved by AICTE | Affiliated to Osmania University | Estd. 2003.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Line 4: NAAC
    p4 = text_cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Twips(0)
    p4.paragraph_format.space_after = Twips(0)
    run = p4.add_run("Accredited with 'A' grade by NAAC | Accredited by NBA")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    # Line 5: Department
    p5 = text_cell.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p5.paragraph_format.space_before = Twips(0)
    p5.paragraph_format.space_after = Twips(0)
    run = p5.add_run(f"Department of {department_name}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True

def get_student_complete_data(student_roll, subjects_data, backlog_data=None):
    """Get complete data for a student across all subjects.
    Student name and father name are retrieved exclusively from Student Info file (backlog_data).
    If student is not found in Student Info, these fields will be empty strings."""
    student_complete_data = {
        'personal_info': {},
        'subjects': []
    }
    
    # Get student_name and father_name EXCLUSIVELY from Student Info file (backlog_data)
    father_name = ''
    student_name_from_backlog = ''
    if backlog_data is not None:
        # Find roll column in backlog data
        roll_col = None
        for col in ['roll_no', 'roll no', 'rollno']:
            if col in backlog_data.columns:
                roll_col = col
                break
        if roll_col:
            student_roll_str = str(student_roll).strip()
            student_backlog = backlog_data[backlog_data[roll_col].astype(str).str.strip() == student_roll_str]
            if not student_backlog.empty:
                # Get father_name from Student Info
                for col in ['father_name', 'father name', 'fathername']:
                    if col in student_backlog.columns:
                        val = student_backlog.iloc[0][col]
                        if pd.notna(val) and str(val).strip():
                            father_name = str(val).strip()
                            break
                # Get student_name from Student Info
                for col in ['student_name', 'student name', 'name']:
                    if col in student_backlog.columns:
                        val = student_backlog.iloc[0][col]
                        if pd.notna(val) and str(val).strip():
                            student_name_from_backlog = str(val).strip()
                            break
    
    student_roll_str = str(student_roll).strip()
    for subject_name, subject_df in subjects_data.items():
        student_data = subject_df[subject_df['roll_no'].astype(str).str.strip() == student_roll_str]
        if not student_data.empty:
            student_info = student_data.iloc[0].to_dict()
            if not student_complete_data['personal_info']:
                student_complete_data['personal_info'] = {
                    'roll_no': student_info['roll_no'],
                    'student_name': student_name_from_backlog if student_name_from_backlog else f"Student {student_roll}",
                    'father_name': father_name
                }
            subject_data = {
                'subject_name': subject_name,
                'dt_marks': student_info.get('dt_marks', 0),
                'aat_marks': student_info.get('aat_marks', 0),
                'at_marks': student_info.get('at_marks', 0),
                'total_marks': student_info.get('total_marks', 0),
                'dtde_marks': student_info.get('dtde_marks', 0),
                'cie_marks': student_info.get('cie_marks', 0),
                'attendance_conducted': student_info.get('attendance_conducted', 0),
                'attendance_present': student_info.get('attendance_present', 0),
                'is_lab': student_info.get('is_lab', False),
                'has_original_lab_marks': student_info.get('has_original_lab_marks', False)
            }
            student_complete_data['subjects'].append(subject_data)
    return student_complete_data


def _safe_float(val):
    """Safely convert a value to float, treating NaN/None/empty as 0."""
    import math
    if val is None:
        return 0.0
    if isinstance(val, float) and math.isnan(val):
        return 0.0
    if isinstance(val, str):
        val = val.strip()
        if val == '' or val.lower() == 'ab':
            return 0.0
    try:
        result = float(val)
        if math.isnan(result):
            return 0.0
        return result
    except (ValueError, TypeError):
        return 0.0


def _format_mark_value(val):
    """Format a mark value for display: AB shows as 'AB', numbers show rounded."""
    import math
    if isinstance(val, str) and val.strip().lower() == 'ab':
        return 'AB'
    try:
        f = float(val)
        if math.isnan(f):
            return '0'
        return str(round(f))
    except (ValueError, TypeError):
        return str(val) if val not in ('', None) else ''


def _style_cell_text(cell, font_name='Times New Roman', font_size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Apply consistent styling to all runs in a cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
        paragraph.alignment = alignment


def _add_student_report_content(doc, student_complete_data, department_name, report_date, academic_year, semester, attendance_start="", attendance_end="", template="Detailed", include_backlog=True, include_notes=True, backlog_data=None):
    """Add a single student's report content to the document. 
    Shared logic between individual and consolidated report generation."""
    from docx.shared import Twips, Emu
    
    add_logo_and_header(doc, department_name)
    
    # Date line - right aligned
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(f"Date: {report_date}")
    date_run.font.name = 'Times New Roman'
    date_run.font.size = Pt(12)
    date_run.font.bold = True
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Progress Report")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.underline = True
    
    personal_info = student_complete_data['personal_info']
    
    # Academic Year line with semester right-aligned
    academic_para = doc.add_paragraph()
    academic_para.paragraph_format.space_after = Pt(0)
    academic_run = academic_para.add_run(f"Academic Year          : {academic_year}")
    academic_run.font.name = 'Times New Roman'
    academic_run.font.size = Pt(12)
    academic_run.font.bold = True
    academic_para.add_run("\t\t\t\t\t\t\t\t")
    semester_run = academic_para.add_run(f"{semester}")
    semester_run.font.name = 'Times New Roman'
    semester_run.font.size = Pt(12)
    semester_run.font.bold = True
    
    # Roll No line - bold, adjusted alignment
    roll_para = doc.add_paragraph()
    roll_para.paragraph_format.space_before = Pt(0)
    roll_para.paragraph_format.space_after = Pt(0)
    roll_run = roll_para.add_run(f"Roll No.                      : {personal_info['roll_no']}")
    roll_run.font.name = 'Times New Roman'
    roll_run.font.size = Pt(12)
    roll_run.font.bold = True
    
    # Student Name line - bold, adjusted alignment
    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(0)
    name_run = name_para.add_run(f"Name of the Student : {str(personal_info['student_name']).upper()}")
    name_run.font.name = 'Times New Roman'
    name_run.font.size = Pt(12)
    name_run.font.bold = True
    
    # Father Name line - bold, adjusted alignment
    father_para = doc.add_paragraph()
    father_para.paragraph_format.space_before = Pt(0)
    father_para.paragraph_format.space_after = Pt(0)
    father_run = father_para.add_run(f"Name of the Father   : {str(personal_info['father_name']).upper()}")
    father_run.font.name = 'Times New Roman'
    father_run.font.size = Pt(12)
    father_run.font.bold = True
    
    if template == "Detailed":
        greeting_para = doc.add_paragraph()
        greeting_para.paragraph_format.space_before = Pt(0)
        greeting_para.paragraph_format.space_after = Pt(0)
        greeting_run = greeting_para.add_run("Dear Parent/Guardian,")
        greeting_run.font.name = 'Times New Roman'
        greeting_run.font.size = Pt(12)
        greeting_run.font.bold = True
        
        desc_para = doc.add_paragraph()
        desc_para.paragraph_format.space_before = Pt(0)
        desc_run = desc_para.add_run("The following are the details of the attendance and Continuous Internal Evaluation-1 of your ward. It is furnished for your information.")
        desc_run.font.name = 'Times New Roman'
        desc_run.font.size = Pt(12)

    subjects = student_complete_data['subjects']
    if not subjects:
        return
    
    # Sort: theory first, then labs
    theory_subjects = [s for s in subjects if not s.get('is_lab', False)]
    lab_subjects = [s for s in subjects if s.get('is_lab', False)]
    subjects_sorted = theory_subjects + lab_subjects
    
    # ======== BUILD THE TABLE ========
    # Table has 8 columns: S.No | Course Title | Conducted | Attended | Col5 | Col6 | Col7 | Col8
    table = doc.add_table(rows=2, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    top = table.rows[0].cells
    bottom = table.rows[1].cells

    # Merge S.No. and Course Title vertically
    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(1, 1))
    top[0].text = 'S. No.'
    top[1].text = 'Course Title'

    # Attendance group spanning two columns
    attendance_top = table.cell(0, 2)
    attendance_top.merge(table.cell(0, 3))
    period_text = 'Attendance'
    if attendance_start and attendance_end:
        period_text += f"\n(From {attendance_start} to {attendance_end})"
    attendance_top.text = period_text
    bottom[2].text = 'No. of Classes\nConducted'
    bottom[3].text = 'No. of Classes\nAttended'

    # CIE-1 Marks group spanning four columns
    marks_top = table.cell(0, 4)
    marks_top.merge(table.cell(0, 7))
    marks_top.text = 'CIE-1 Marks'
    bottom[4].text = 'DT\n(20)'
    bottom[5].text = 'AT-1\n(10)'
    bottom[6].text = 'AAT-1\n(10)'
    bottom[7].text = 'Total\n(40)'

    # Style header rows
    for row in [table.rows[0], table.rows[1]]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.bold = True
                    run.font.size = Pt(12)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ======== DATA ROWS ========
    total_attendance_conducted = 0
    total_attendance_present = 0
    total_marks_sum = 0
    total_max_marks = 0
    num_lab_subjects_with_marks = 0
    serial_num = 0
    
    lab_section_header_added = False
    
    for subject in subjects_sorted:
        attendance_conducted = subject['attendance_conducted']
        attendance_present = subject['attendance_present']
        total_attendance_conducted += attendance_conducted
        total_attendance_present += attendance_present
        is_lab = bool(subject.get('is_lab', False))
        
        if is_lab and not lab_section_header_added:
            # Add "Lab Details" sub-header row before first lab subject
            lab_header_row = table.add_row()
            lab_hdr_cells = lab_header_row.cells
            # Merge S.No + Course Title for "Lab Details" label
            merged_lab_label = lab_hdr_cells[0].merge(lab_hdr_cells[1])
            merged_lab_label.merge(lab_hdr_cells[2])
            merged_lab_label = table.cell(len(table.rows) - 1, 0)
            merged_lab_label.merge(table.cell(len(table.rows) - 1, 3))
            merged_lab_label.text = 'Lab Details'
            # DTDE (15) header - merge cols 4+5
            dtde_header = table.cell(len(table.rows) - 1, 4)
            dtde_header.merge(table.cell(len(table.rows) - 1, 5))
            dtde_header.text = 'DTDE (15)'
            # CIE (10) header
            cie_header = table.cell(len(table.rows) - 1, 6)
            cie_header.text = 'CIE\n(10)'
            # Total (25) header
            total_header = table.cell(len(table.rows) - 1, 7)
            total_header.text = 'Total (25)'
            
            # Style the lab header row
            for cell in lab_header_row.cells:
                _style_cell_text(cell, bold=True)
            lab_section_header_added = True
        
        serial_num += 1
        data_row = table.add_row()
        data_cells = data_row.cells
        
        if is_lab:
            # Lab subject row
            has_orig_lab_marks = bool(subject.get('has_original_lab_marks', False))
            
            dtde_val = subject.get('dtde_marks', 0)
            cie_val = subject.get('cie_marks', 0)
            
            # Check for AB values
            dtde_is_ab = isinstance(dtde_val, str) and str(dtde_val).strip().lower() == 'ab'
            cie_is_ab = isinstance(cie_val, str) and str(cie_val).strip().lower() == 'ab'
            
            dtde_numeric = 0 if dtde_is_ab else _safe_float(dtde_val)
            cie_numeric = 0 if cie_is_ab else _safe_float(cie_val)
            
            lab_total = dtde_numeric + cie_numeric
            
            if has_orig_lab_marks:
                num_lab_subjects_with_marks += 1
                total_marks_sum += lab_total
            
            # Merge S.No cell
            data_cells[0].text = str(serial_num)
            data_cells[1].text = str(subject['subject_name']).title()
            data_cells[2].text = str(attendance_conducted)
            data_cells[3].text = str(attendance_present)
            
            # DTDE in merged cols 4+5
            merged_dtde = data_cells[4].merge(data_cells[5])
            if has_orig_lab_marks:
                merged_dtde.text = 'AB' if dtde_is_ab else str(round(dtde_numeric))
            else:
                merged_dtde.text = '-'
            
            # CIE in col 6
            if has_orig_lab_marks:
                data_cells[6].text = 'AB' if cie_is_ab else str(round(cie_numeric))
            else:
                data_cells[6].text = '-'
            
            # Total in col 7
            if has_orig_lab_marks:
                data_cells[7].text = str(round(lab_total))
            else:
                data_cells[7].text = '-'
            
            # Style all cells
            for i, cell in enumerate(data_row.cells):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                    if i == 1:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Theory subject row
            dt_val = subject['dt_marks']
            at_val = subject['at_marks']
            aat_val = subject['aat_marks']
            
            dt_is_ab = isinstance(dt_val, str) and str(dt_val).strip().lower() == 'ab'
            at_is_ab = isinstance(at_val, str) and str(at_val).strip().lower() == 'ab'
            aat_is_ab = isinstance(aat_val, str) and str(aat_val).strip().lower() == 'ab'
            
            dt_numeric = 0 if dt_is_ab else _safe_float(dt_val)
            at_numeric = 0 if at_is_ab else _safe_float(at_val)
            aat_numeric = 0 if aat_is_ab else _safe_float(aat_val)
            
            dt_marks = 'AB' if dt_is_ab else dt_numeric
            at_marks = 'AB' if at_is_ab else at_numeric
            aat_marks = 'AB' if aat_is_ab else aat_numeric
            total_marks = dt_numeric + at_numeric + aat_numeric
            total_marks_sum += total_marks
            total_max_marks += 40
            
            row_data = [
                str(serial_num),
                str(subject['subject_name']).title(),
                str(attendance_conducted),
                str(attendance_present),
                _format_mark_value(dt_marks),
                _format_mark_value(at_marks),
                _format_mark_value(aat_marks),
                str(round(total_marks))
            ]
            for i, data in enumerate(row_data):
                if i < len(data_cells):
                    data_cells[i].text = data
                    for paragraph in data_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)
                        if i == 1:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add lab subjects with marks to max marks
    total_max_marks += num_lab_subjects_with_marks * 25
    
    # ======== TOTAL ROW ========
    total_row = table.add_row()
    total_cells = total_row.cells
    total_row_idx = len(table.rows) - 1
    merged_sno_title = table.cell(total_row_idx, 0).merge(table.cell(total_row_idx, 1))
    merged_sno_title.text = "TOTAL"
    for p in merged_sno_title.paragraphs:
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total_cells[2].text = str(total_attendance_conducted)
    total_cells[3].text = str(total_attendance_present)
    try:
        merged_total = total_cells[4].merge(total_cells[5])
        merged_total = merged_total.merge(total_cells[6])
        merged_total = merged_total.merge(total_cells[7])
        if total_max_marks > 0:
            merged_total.text = str(round(total_marks_sum))
        else:
            merged_total.text = ''
    except Exception:
        pass
    
    overall_attendance_percent = (total_attendance_present / total_attendance_conducted * 100) if total_attendance_conducted > 0 else 0
    
    # ======== PERCENTAGE ROW ========
    percent_row = table.add_row()
    percent_cells = percent_row.cells
    percent_sno_idx = len(table.rows) - 1
    merged_sno_pct = table.cell(percent_sno_idx, 0).merge(table.cell(percent_sno_idx, 1))
    merged_sno_pct.text = "Percentage"
    for p in merged_sno_pct.paragraphs:
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    percent_row_idx = len(table.rows) - 1
    merged_attendance_cell = table.cell(percent_row_idx, 2)
    merged_attendance_cell.merge(table.cell(percent_row_idx, 3))
    merged_attendance_cell.text = f"{overall_attendance_percent:.2f}%"
    try:
        merged_marks_cell = table.cell(percent_row_idx, 4).merge(table.cell(percent_row_idx, 5))
        merged_marks_cell = merged_marks_cell.merge(table.cell(percent_row_idx, 6))
        merged_marks_cell = merged_marks_cell.merge(table.cell(percent_row_idx, 7))
        if total_max_marks > 0:
            marks_percentage = (total_marks_sum / total_max_marks) * 100
            merged_marks_cell.text = f"{marks_percentage:.2f}%"
        else:
            merged_marks_cell.text = '-'
        for p in merged_marks_cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass
    
    for row in [total_row, percent_row]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
                    run.font.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Enforce exact widths from reference docx
    try:
        table.autofit = False
        col_widths = [
            Emu(537210),    # S.No.
            Emu(1758315),   # Course Title
            Emu(1080135),   # Attendance Conducted
            Emu(1170305),   # Attendance Attended
            Emu(629920),    # DT (20)
            Emu(540385),    # AT-1 (10)
            Emu(671830),    # AAT-1 (10)
            Emu(560070),    # Total (40)
        ]
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = col_widths[i]
    except Exception:
        pass
    
    if template == "Detailed":
        # Legend
        legend_para = doc.add_paragraph()
        legend_para.paragraph_format.space_before = Pt(0)
        legend_para.paragraph_format.space_after = Pt(0)
        legend_text = "*DT – Descriptive Test   AT-Assignment   AAT- Alternative Assessment Test"
        legend_run = legend_para.add_run(legend_text)
        legend_run.font.name = 'Times New Roman'
        legend_run.font.size = Pt(12)
        
        # Attendance note
        attendance_note = doc.add_paragraph()
        attendance_note.paragraph_format.space_before = Pt(0)
        attendance_note.paragraph_format.space_after = Pt(0)
        attendance_status = "Poor" if overall_attendance_percent < 75 else "Satisfactory"
        note_text = f"Your ward's attendance is {overall_attendance_percent:.2f}% which is {attendance_status}."
        note_run = attendance_note.add_run(note_text)
        note_run.font.name = 'Times New Roman'
        note_run.font.size = Pt(11)
        note_run.font.bold = False
        if overall_attendance_percent < 75:
            note_run.font.color.rgb = RGBColor(255, 0, 0)
        
        if include_notes:
            important_para = doc.add_paragraph()
            important_para.paragraph_format.space_after = Pt(0)
            important_run = important_para.add_run("Important Note:")
            important_run.font.name = 'Times New Roman'
            important_run.font.size = Pt(12)
            important_run.font.bold = True
            
            notes_para = doc.add_paragraph()
            notes_para.paragraph_format.space_after = Pt(0)
            
            run1 = notes_para.add_run("     ➤ As per the ")
            run1.font.name = 'Times New Roman'
            run1.font.size = Pt(11)
            
            run2 = notes_para.add_run("Osmania University")
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(11)
            run2.font.bold = True
            
            run3 = notes_para.add_run(" rules, a student must have minimum attendance of 75% in aggregate of all the subjects to be eligible or promoted for the next year. Students having less than 75% attendance in aggregate will not be issued Hall Ticket for the examination, such students will come under Condonation/Detention category.\n")
            run3.font.name = 'Times New Roman'
            run3.font.size = Pt(11)
            
            run4 = notes_para.add_run("     ➤ As per State Government rules, the student is not eligible for Scholarship if the attendance is less than 75%.")
            run4.font.name = 'Times New Roman'
            run4.font.size = Pt(11)
            run4.font.bold = True
        
        if include_backlog:
            backlog_para = doc.add_paragraph()
            backlog_para.paragraph_format.space_after = Pt(0)
            backlog_run = backlog_para.add_run("Backlog Data:")
            backlog_run.font.name = 'Times New Roman'
            backlog_run.font.size = Pt(12)
            backlog_run.font.bold = True
            
            import re
            semester_map = {'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5, 'VI': 6, 'VII': 7, 'VIII': 8}
            semester_num = 4
            semester_match = re.search(r'\b(VIII|VII|VI|V|IV|III|II|I)\b', semester)
            if semester_match:
                semester_num = semester_map.get(semester_match.group(1), 4)
            
            roman_numerals = {1: 'I', 2: 'II', 3: 'III', 4: 'IV', 5: 'V', 6: 'VI', 7: 'VII', 8: 'VIII'}
            num_prev_semesters = max(1, semester_num - 1)
            backlog_headers = [f'{roman_numerals.get(i, str(i))} Sem.' for i in range(1, num_prev_semesters + 1)]
            backlog_headers.append('Remarks by Head of the Department')
            
            num_cols = len(backlog_headers)
            backlog_table = doc.add_table(rows=2, cols=num_cols)
            backlog_table.style = 'Table Grid'
            backlog_table.autofit = False
            
            # Use dynamic widths but preserve the large Remarks column
            for row in backlog_table.rows:
                for i, cell in enumerate(row.cells):
                    if i == num_cols - 1:
                        cell.width = Emu(2708275) # Remarks column exact width
                    else:
                        cell.width = Emu(800000) # Base width for semester columns
            
            backlog_hdr_cells = backlog_table.rows[0].cells
            for i, header in enumerate(backlog_headers):
                backlog_hdr_cells[i].text = header
                _style_cell_text(backlog_hdr_cells[i], bold=True)
            
            backlog_data_cells = backlog_table.rows[1].cells
            
            student_roll = student_complete_data['personal_info']['roll_no']
            student_backlog = None
            if backlog_data is not None:
                roll_col = None
                for col in ['roll_no', 'roll no', 'rollno']:
                    if col in backlog_data.columns:
                        roll_col = col
                        break
                if roll_col:
                    student_backlog_row = backlog_data[backlog_data[roll_col].astype(str).str.strip() == str(student_roll).strip()]
                    if not student_backlog_row.empty:
                        student_backlog = student_backlog_row.iloc[0]
            
            for i, cell in enumerate(backlog_data_cells):
                if i < num_prev_semesters:
                    sem_num = i + 1
                    possible_col_names = [f'sem {sem_num}', f'sem{sem_num}', f'Sem {sem_num}', f'Sem{sem_num}']
                    cell_value = None
                    if student_backlog is not None:
                        for col_name in possible_col_names:
                            if col_name in student_backlog.index:
                                cell_value = student_backlog[col_name]
                                break
                    if cell_value is not None and pd.notna(cell_value):
                        cell.text = str(cell_value)
                    else:
                        cell.text = '-'
                else:
                    total_backlogs = 0
                    if student_backlog is not None:
                        for sem_i in range(1, num_prev_semesters + 1):
                            for col_name in [f'sem {sem_i}', f'sem{sem_i}', f'Sem {sem_i}', f'Sem{sem_i}']:
                                if col_name in student_backlog.index:
                                    val = student_backlog[col_name]
                                    if pd.notna(val):
                                        try:
                                            total_backlogs += int(val)
                                        except (ValueError, TypeError):
                                            pass
                                    break
                    cie_percent = (total_marks_sum / total_max_marks * 100) if total_max_marks > 0 else 0
                    remark = generate_hod_remark(overall_attendance_percent, cie_percent, total_backlogs)
                    cell.text = remark
                _style_cell_text(cell)
        
        doc.add_paragraph()
        # Signature line
        signature_para = doc.add_paragraph()
        signature_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        signature_para.paragraph_format.space_before = Pt(6)
        signature_para.paragraph_format.space_after = Pt(6)
        signature_text = "Sign. of the student: _______________________Sign. of the Parent/Guardian: _________________________"
        signature_run = signature_para.add_run(signature_text)
        signature_run.font.name = 'Times New Roman'
        signature_run.font.size = Pt(11)
        signature_run.font.bold = True


def create_comprehensive_student_report(student_complete_data, department_name, report_date, academic_year, semester, attendance_start="", attendance_end="", template="Detailed", include_backlog=True, include_notes=True, backlog_data=None):
    """Create a comprehensive Word document report for a student with customizable template"""
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.1)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    _add_student_report_content(doc, student_complete_data, department_name, report_date, academic_year, semester, attendance_start, attendance_end, template, include_backlog, include_notes, backlog_data)
    
    return doc

def generate_student_reports(student_roll, subjects_data, department_name, report_date, academic_year, semester, attendance_start="", attendance_end="", template="Detailed", include_backlog=True, include_notes=True, backlog_data=None):
    """Generate a comprehensive report for a single student in Word format"""
    student_complete_data = get_student_complete_data(student_roll, subjects_data, backlog_data)
    if not student_complete_data['subjects']:
        return {}
    doc = create_comprehensive_student_report(student_complete_data, department_name, report_date, academic_year, semester, attendance_start, attendance_end, template, include_backlog, include_notes, backlog_data)
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    student_name = student_complete_data['personal_info']['student_name']
    return {
        f"{student_name}_Comprehensive_Report_docx": doc_buffer.getvalue()
    }

def create_consolidated_all_students_report(all_students_data, subjects_data, department_name, report_date, academic_year, semester, attendance_start="", attendance_end="", template="Detailed", include_backlog=True, include_notes=True, backlog_data=None):
    """Create a single Word document containing all student reports, each on a separate page"""
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.1)
        section.bottom_margin = Inches(0.45)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    for idx, student_roll in enumerate(all_students_data):
        if idx > 0:
            doc.add_page_break()
        student_complete_data = get_student_complete_data(student_roll, subjects_data, backlog_data)
        if student_complete_data['subjects']:
            _add_student_report_content(doc, student_complete_data, department_name, report_date, academic_year, semester, attendance_start, attendance_end, template, include_backlog, include_notes, backlog_data)
    
    return doc

def generate_comprehensive_reports(all_students, subjects_data, department_name, report_date, academic_year, semester, attendance_start="", attendance_end="", template="Detailed", include_backlog=True, include_notes=True, backlog_data=None):
    """Generate comprehensive reports for all students in parallel"""
    individual_reports = {}
    with concurrent.futures.ThreadPoolExecutor() as executor:
        future_to_roll = {
            executor.submit(generate_student_reports, roll, subjects_data, department_name, report_date, academic_year, semester, attendance_start, attendance_end, template, include_backlog, include_notes, backlog_data): roll
            for roll in all_students
        }
        for future in concurrent.futures.as_completed(future_to_roll):
            roll = future_to_roll[future]
            try:
                report = future.result()
                if report:
                    student_name = list(report.keys())[0].split('_Comprehensive_Report_docx')[0]
                    individual_reports[roll] = {
                        'docx_content': report[f"{student_name}_Comprehensive_Report_docx"],
                        'student_name': student_name
                    }
            except Exception as e:
                print(f"Error generating report for {roll}: {str(e)}")
    consolidated_doc = create_consolidated_all_students_report(all_students, subjects_data, department_name, report_date, academic_year, semester, attendance_start, attendance_end, template, include_backlog, include_notes, backlog_data)
    consolidated_buffer = BytesIO()
    consolidated_doc.save(consolidated_buffer)
    consolidated_buffer.seek(0)
    return individual_reports, consolidated_buffer.getvalue()
