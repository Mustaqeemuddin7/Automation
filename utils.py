# utils.py
# Utility functions for the LORDS Institute Progress Report System

import re
import pandas as pd
import streamlit as st
from io import BytesIO
import mammoth
from config import COLUMN_MAPPINGS

def normalize_column_name(col_name):
    """Normalize column name by converting to lowercase and removing spaces/underscores"""
    if not isinstance(col_name, str):
        return ""
    normalized = re.sub(r'[\s_\.\-]', '', col_name.lower())
    return normalized

def map_column_name(col_name):
    """Map a column name to the standardized name based on variations"""
    normalized = normalize_column_name(col_name)
    for standard_name, variations in COLUMN_MAPPINGS.items():
        if normalized in [normalize_column_name(v) for v in variations]:
            return standard_name
    return col_name

@st.cache_data
def process_subject_files(uploaded_files):
    """Process multiple Excel files (theory and lab), each representing a subject.
    Labs may contain only attendance columns; theory files include marks.
    """
    try:
        subjects_data = {}
        all_students = set()
        # Minimal required columns for any subject (lab or theory)
        # student_name and father_name now come from Student Info file
        minimal_required = ['roll_no', 'attendance_conducted', 'attendance_present']
        for file in uploaded_files:
            subject_name = file.name.split('.')[0]
            df = pd.read_excel(file)
            column_mapping = {}
            for col in df.columns:
                standardized_col = map_column_name(col)
                if standardized_col in COLUMN_MAPPINGS.keys():
                    column_mapping[col] = standardized_col
            df = df.rename(columns=column_mapping)
            # Validate minimal columns
            missing_min = [col for col in minimal_required if col not in df.columns]
            if missing_min:
                return None, None, f"Missing required columns in {subject_name}: {', '.join(missing_min)}"

            # Determine if this is a lab file (no theory marks present)
            has_dt = 'dt_marks' in df.columns
            has_st = 'st_marks' in df.columns
            has_at = 'at_marks' in df.columns
            has_lab_marks = 'lab_marks' in df.columns
            is_lab_file = not (has_dt or has_st or has_at)

            # Ensure marks columns exist; for labs, fill with 0 to avoid UI errors
            # But preserve 'ab'/'AB' string values in existing marks columns
            for col in ['dt_marks', 'st_marks', 'at_marks', 'total_marks']:
                if col not in df.columns:
                    df[col] = 0
            if not has_lab_marks:
                df['lab_marks'] = 0
            # Convert marks columns to object dtype so mixed types (int/float + 'ab')
            # don't cause Arrow serialization errors in Streamlit
            for col in ['dt_marks', 'st_marks', 'at_marks', 'total_marks', 'lab_marks']:
                if col in df.columns:
                    df[col] = df[col].astype(object)
            df['is_lab'] = is_lab_file
            subjects_data[subject_name] = df
            if 'roll_no' in df.columns:
                all_students.update(df['roll_no'].tolist())
        return subjects_data, list(all_students), None
    except Exception as e:
        return None, None, str(e)

def preview_report(doc_content, title):
    """Generate HTML and text previews for a report"""
    mammoth_style_map = """
    p[style-name='Normal'] => p
    table => table.table-bordered
    """
    mammoth_custom_styles = """
    <style>
        .table-bordered {
            border-collapse: collapse;
            width: 100%;
            margin: 10px 0;
        }
        .table-bordered th, .table-bordered td {
            border: 1px solid #ddd;
            padding: 8px;
            text-align: center;
        }
        .table-bordered th {
            background-color: #f2f2f2;
            font-weight: bold;
        }
    </style>
    """
    with st.expander(f"Preview: {title}"):
        try:
            # HTML Preview
            doc_buffer = BytesIO(doc_content)
            result = mammoth.convert_to_html(doc_buffer, style_map=mammoth_style_map)
            st.markdown(mammoth_custom_styles + result.value, unsafe_allow_html=True)
            if result.messages:
                st.warning("Some formatting may not be preserved in HTML preview.")
        except Exception as e:
            st.warning(f"HTML preview error: {str(e)}. Showing text preview.")
            from docx import Document
            doc = Document(BytesIO(doc_content))
            full_text = [para.text for para in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    full_text.append(" | ".join(cell.text for cell in row.cells))
            st.text_area("Text Preview", "\n".join(full_text), height=400)