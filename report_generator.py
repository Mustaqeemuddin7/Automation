from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
import concurrent.futures
import streamlit as st
from docx.enum.text import WD_BREAK

def add_logo_and_header(doc, department_name):
    """Add institutional header to document with correct font and sizes"""
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("LORDS INSTITUTE OF ENGINEERING & TECHNOLOGY\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    run = header_para.add_run("(Autonomous)\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run = header_para.add_run("Approved by AICTE | Affiliated to Osmania University | Estd. 2003\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run = header_para.add_run("Accredited with 'A' grade by NAAC | Accredited by NBA\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run = header_para.add_run(f"DEPARTMENT OF {department_name.upper()}\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True

def get_student_complete_data(student_roll, subjects_data):
    """Get complete data for a student across all subjects"""
    student_complete_data = {
        'personal_info': {},
        'subjects': []
    }
    for subject_name, subject_df in subjects_data.items():
        student_data = subject_df[subject_df['roll_no'] == student_roll]
        if not student_data.empty:
            student_info = student_data.iloc[0].to_dict()
            if not student_complete_data['personal_info']:
                student_complete_data['personal_info'] = {
                    'roll_no': student_info['roll_no'],
                    'student_name': student_info['student_name'],
                    'father_name': student_info['father_name']
                }
            subject_data = {
                'subject_name': subject_name,
                'dt_marks': student_info.get('dt_marks', 0),
                'st_marks': student_info.get('st_marks', 0),
                'at_marks': student_info.get('at_marks', 0),
                'total_marks': student_info.get('total_marks', 0),
                'attendance_conducted': student_info.get('attendance_conducted', 0),
                'attendance_present': student_info.get('attendance_present', 0),
                'is_lab': student_info.get('is_lab', False)
            }
            student_complete_data['subjects'].append(subject_data)
    return student_complete_data

def create_comprehensive_student_report(student_complete_data, department_name, report_date, academic_year, semester, attendance_start="", attendance_end="", template="Detailed", include_backlog=True, include_notes=True):
    """Create a comprehensive Word document report for a student with customizable template"""
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    add_logo_and_header(doc, department_name)
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(f"Date: {report_date}")
    date_run.font.name = 'Times New Roman'
    date_run.font.size = Pt(10)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Progress Report")
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.underline = True
    personal_info = student_complete_data['personal_info']
    details_para = doc.add_paragraph()
    details_text = f"Academic Year: {academic_year}".ljust(60) + f"{semester}\n"
    details_text += f"Roll No.              : {personal_info['roll_no']}\n"
    details_text += f"Name of the Student : {personal_info['student_name']}\n"
    details_text += f"Name of the Father   : {personal_info['father_name']}\n"
    if template == "Detailed":
        details_text += "Dear Parent/Guardian,\n\n"
        details_text += "The following are the details of the attendance and Continuous Internal Evaluation-1 of your ward. It is furnished for your information."
    details_run = details_para.add_run(details_text)
    details_run.font.name = 'Times New Roman'
    details_run.font.size = Pt(10)
    subjects = student_complete_data['subjects']
    if subjects:
        # Create a two-row header with grouped columns like the reference image
        table = doc.add_table(rows=2, cols=8)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Top header row labels
        top = table.rows[0].cells
        bottom = table.rows[1].cells

        # Merge S.No. and Course Title vertically
        table.cell(0, 0).merge(table.cell(1, 0))
        table.cell(0, 1).merge(table.cell(1, 1))
        top[0].text = 'S. No.'
        top[1].text = 'Course Title'

        # Attendance group spanning two columns
        attendance_top = table.cell(0, 2)
        attendance_top.merge(table.cell(0, 3))
        period_text = 'Attendance'
        if attendance_start and attendance_end:
            period_text += f"\n(From {attendance_start} to {attendance_end})"
        attendance_top.text = period_text
        bottom[2].text = 'No. of Classes\nConducted'
        bottom[3].text = 'No. of Classes\nAttended'

        # CIE-1 Marks group spanning four columns
        marks_top = table.cell(0, 4)
        marks_top.merge(table.cell(0, 7))
        marks_top.text = 'CIE-1 Marks'
        bottom[4].text = 'DT\n(20)'
        bottom[5].text = 'ST\n(10)'
        bottom[6].text = 'AT\n(10)'
        bottom[7].text = 'Total\n(40)'

        # Style header rows
        for row in [table.rows[0], table.rows[1]]:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.bold = True
                        run.font.size = Pt(9)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # We'll enforce S.No. width after the table is fully built
        # Use default auto-fit widths (restored)
        total_attendance_conducted = 0
        total_attendance_present = 0
        total_marks_sum = 0
        # Sort: theory first (is_lab False/absent), then labs
        subjects_sorted = sorted(subjects, key=lambda s: 1 if s.get('is_lab', False) else 0)
        for idx, subject in enumerate(subjects_sorted):
            data_row = table.add_row()
            data_cells = data_row.cells
            attendance_conducted = subject['attendance_conducted']
            attendance_present = subject['attendance_present']
            total_attendance_conducted += attendance_conducted
            total_attendance_present += attendance_present
            is_lab = bool(subject.get('is_lab', False))
            if is_lab:
                # For labs, we will merge DT/ST/AT/Total cells later and set a single '-'
                dt_marks = ''
                st_marks = ''
                at_marks = ''
                total_marks = ''
            else:
                dt_marks = subject['dt_marks'] * 20 / 30
                st_marks = subject['st_marks']
                at_marks = subject['at_marks']
                total_marks = dt_marks + st_marks + at_marks
                total_marks_sum += total_marks
            row_data = [
                str(idx + 1),
                subject['subject_name'],
                str(attendance_conducted),
                str(attendance_present),
                (str(round(dt_marks)) if dt_marks not in ('', '-') else ''),
                (str(st_marks) if st_marks not in ('', '-') else ''),
                (str(at_marks) if at_marks not in ('', '-') else ''),
                (str(round(total_marks)) if total_marks not in ('', '-') else '')
            ]
            for i, data in enumerate(row_data):
                if i < len(data_cells):
                    data_cells[i].text = data
                    for paragraph in data_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(9)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # For lab subjects, merge DT, ST, AT and Total cells into one and set a single '-'
            if is_lab:
                try:
                    merged = data_cells[4].merge(data_cells[5])
                    merged = merged.merge(data_cells[6])
                    merged = merged.merge(data_cells[7])
                    # Set a single hyphen and center it
                    merged.text = '-'
                    for paragraph in merged.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(9)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
        total_row = table.add_row()
        total_cells = total_row.cells
        total_cells[1].text = "TOTAL"
        total_cells[2].text = str(total_attendance_conducted)
        total_cells[3].text = str(total_attendance_present)
        # Merge DT, ST, AT, Total cells in TOTAL row
        try:
            merged_total = total_cells[4].merge(total_cells[5])
            merged_total = merged_total.merge(total_cells[6])
            merged_total = merged_total.merge(total_cells[7])
            merged_total.text = ''
        except Exception:
            pass
        overall_attendance_percent = (total_attendance_present / total_attendance_conducted * 100) if total_attendance_conducted > 0 else 0
        percent_row = table.add_row()
        percent_cells = percent_row.cells
        percent_cells[1].text = "Percentage"
        # Merge attendance columns (Conducted and Attended) into one cell for percentage value
        percent_row_idx = len(table.rows) - 1
        merged_attendance_cell = table.cell(percent_row_idx, 2)
        merged_attendance_cell.merge(table.cell(percent_row_idx, 3))
        merged_attendance_cell.text = f"{overall_attendance_percent:.2f}%"
        # Merge DT/ST/AT/Total for percentage row and place '-'
        try:
            merged_marks_cell = table.cell(percent_row_idx, 4).merge(table.cell(percent_row_idx, 5))
            merged_marks_cell = merged_marks_cell.merge(table.cell(percent_row_idx, 6))
            merged_marks_cell = merged_marks_cell.merge(table.cell(percent_row_idx, 7))
            merged_marks_cell.text = '-'
            for p in merged_marks_cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass
        for row in [total_row, percent_row]:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(9)
                        run.font.bold = True
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Enforce compact widths after table is complete to match reference image
        try:
            table.autofit = False
            # Set specific column widths as per specifications
            table.columns[0].width = Inches(0.1)  # S.No. = 0.1 cm
            table.columns[1].width = Inches(0.3)  # Course Title = 3 times S.No. (0.3 cm)
            table.columns[2].width = Inches(0.5)  # Attendance = 5 times S.No. (0.5 cm)
            table.columns[3].width = Inches(0.5)  # Attendance = 5 times S.No. (0.5 cm)
            # CIE-1 Marks columns = 5 times S.No. (0.5 cm)
            table.columns[4].width = Inches(0.5)  # DT (20)
            table.columns[5].width = Inches(0.5)  # ST (10)
            table.columns[6].width = Inches(0.5)  # AT (10)
            table.columns[7].width = Inches(0.5)  # Total (40)
        except Exception:
            # Fallback for individual cell width setting
            for r in table.rows:
                r.cells[0].width = Inches(0.1)  # S.No. = 0.1 cm
                r.cells[1].width = Inches(0.3)  # Course Title = 3 times S.No. (0.3 cm)
                r.cells[2].width = Inches(0.5)  # Attendance = 5 times S.No. (0.5 cm)
                r.cells[3].width = Inches(0.5)  # Attendance = 5 times S.No. (0.5 cm)
                r.cells[4].width = Inches(0.5)  # CIE-1 Marks = 5 times S.No. (0.5 cm)
                r.cells[5].width = Inches(0.5)  # CIE-1 Marks = 5 times S.No. (0.5 cm)
                r.cells[6].width = Inches(0.5)  # CIE-1 Marks = 5 times S.No. (0.5 cm)
                r.cells[7].width = Inches(0.5)  # CIE-1 Marks = 5 times S.No. (0.5 cm)
        if template == "Detailed":
            legend_para = doc.add_paragraph()
            legend_text = "*DT – Descriptive Test  ST-Surprise Test  AT- Assignment"
            legend_run = legend_para.add_run(legend_text)
            legend_run.font.name = 'Times New Roman'
            legend_run.font.size = Pt(9)
            attendance_note = doc.add_paragraph()
            attendance_status = "Poor" if overall_attendance_percent < 75 else "Satisfactory"
            note_text = f"Your ward's attendance is {overall_attendance_percent:.2f}% which is {attendance_status}."
            note_run = attendance_note.add_run(note_text)
            note_run.font.name = 'Times New Roman'
            note_run.font.size = Pt(9)
            note_run.font.bold = True
            if overall_attendance_percent < 75:
                note_run.font.color.rgb = RGBColor(255, 0, 0)
            if include_notes:
                important_para = doc.add_paragraph()
                important_run = important_para.add_run("Important Note:")
                important_run.font.name = 'Times New Roman'
                important_run.font.size = Pt(10)
                important_run.font.bold = True
                notes_text = "→ As per the Osmania University rules, a student must have minimum attendance of 75% in aggregate of all the subjects to be eligible or promoted for the next year. Students having less than 75% attendance in aggregate will not be issued Hall Ticket for the examination, such students will come under Condonation/Detention category.\n"
                notes_text += "→ As per State Government rules, the student is not eligible for Scholarship if the attendance is less than 75%."
                notes_para = doc.add_paragraph()
                notes_run = notes_para.add_run(notes_text)
                notes_run.font.name = 'Times New Roman'
                notes_run.font.size = Pt(9)
            if include_backlog:
                backlog_para = doc.add_paragraph()
                backlog_run = backlog_para.add_run("Backlog Data:")
                backlog_run.font.name = 'Times New Roman'
                backlog_run.font.size = Pt(10)
                backlog_run.font.bold = True
                backlog_table = doc.add_table(rows=2, cols=4)
                backlog_table.style = 'Table Grid'
                backlog_headers = ['I Sem.', 'II Sem.', 'III Sem.', 'Remarks by Head of the Department']
                backlog_hdr_cells = backlog_table.rows[0].cells
                for i, header in enumerate(backlog_headers):
                    backlog_hdr_cells[i].text = header
                    for paragraph in backlog_hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.bold = True
                            run.font.size = Pt(9)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                backlog_data_cells = backlog_table.rows[1].cells
                for cell in backlog_data_cells:
                    cell.text = '-'
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(9)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
                signature_para = doc.add_paragraph()
                signature_text = "Sign. of the student: ____________________   Sign. of the Parent/Guardian: ____________________"
                signature_run = signature_para.add_run(signature_text)
                signature_run.font.name = 'Times New Roman'
                signature_run.font.size = Pt(9)
                # doc.add_paragraph()
                final_signature_para = doc.add_paragraph()
                final_signature_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                mentor_run = final_signature_para.add_run("Mentor")
                mentor_run.font.name = 'Times New Roman'
                mentor_run.font.size = Pt(9)
                spaces = " " * 60
                final_signature_para.add_run(spaces)
                hod_run = final_signature_para.add_run("Head of the Department")
                hod_run.font.name = 'Times New Roman'
                hod_run.font.size = Pt(9)
    return doc

def generate_student_reports(student_roll, subjects_data, department_name, report_date, academic_year, semester, attendance_start="", attendance_end="", template="Detailed", include_backlog=True, include_notes=True):
    """Generate a comprehensive report for a single student in Word format"""
    student_complete_data = get_student_complete_data(student_roll, subjects_data)
    if not student_complete_data['subjects']:
        return {}
    doc = create_comprehensive_student_report(student_complete_data, department_name, report_date, academic_year, semester, attendance_start, attendance_end, template, include_backlog, include_notes)
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    student_name = student_complete_data['personal_info']['student_name']
    return {
        f"{student_name}_Comprehensive_Report_docx": doc_buffer.getvalue()
    }

def create_consolidated_all_students_report(all_students_data, subjects_data, department_name, report_date, academic_year, semester, attendance_start="", attendance_end="", template="Detailed", include_backlog=True, include_notes=True):
    """Create a single Word document containing all student reports, each on a separate page"""
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    for idx, student_roll in enumerate(all_students_data):
        if idx > 0:
            doc.add_page_break()
        student_complete_data = get_student_complete_data(student_roll, subjects_data)
        if student_complete_data['subjects']:
            add_logo_and_header(doc, department_name)
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_run = date_para.add_run(f"Date: {report_date}")
            date_run.font.name = 'Times New Roman'
            date_run.font.size = Pt(10)
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run("Progress Report")
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(14)
            title_run.font.bold = True
            title_run.font.underline = True
            personal_info = student_complete_data['personal_info']
            details_para = doc.add_paragraph()
            details_text = f"Academic Year: {academic_year}".ljust(60) + f"{semester}\n"
            details_text += f"Roll No.              : {personal_info['roll_no']}\n"
            details_text += f"Name of the Student : {personal_info['student_name']}\n"
            details_text += f"Name of the Father   : {personal_info['father_name']}\n"
            if template == "Detailed":
                details_text += "Dear Parent/Guardian,\n\n"
                details_text += "The following are the details of the attendance and Continuous Internal Evaluation-1 of your ward. It is furnished for your information."
            details_run = details_para.add_run(details_text)
            details_run.font.name = 'Times New Roman'
            details_run.font.size = Pt(10)
            subjects = student_complete_data['subjects']
            # Two-row grouped header like the reference image
            table = doc.add_table(rows=2, cols=8)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            top = table.rows[0].cells
            bottom = table.rows[1].cells

            table.cell(0, 0).merge(table.cell(1, 0))
            table.cell(0, 1).merge(table.cell(1, 1))
            top[0].text = 'S. No.'
            top[1].text = 'Course Title'

            attendance_top = table.cell(0, 2)
            attendance_top.merge(table.cell(0, 3))
            period_text = 'Attendance'
            if attendance_start and attendance_end:
                period_text += f"\n(From {attendance_start} to {attendance_end})"
            attendance_top.text = period_text
            bottom[2].text = 'No. of Classes\nConducted'
            bottom[3].text = 'No. of Classes\nAttended'

            marks_top = table.cell(0, 4)
            marks_top.merge(table.cell(0, 7))
            marks_top.text = 'CIE-1 Marks'
            bottom[4].text = 'DT\n(20)'
            bottom[5].text = 'ST\n(10)'
            bottom[6].text = 'AT\n(10)'
            bottom[7].text = 'Total\n(40)'

            for row in [table.rows[0], table.rows[1]]:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.bold = True
                            run.font.size = Pt(9)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Use default auto-fit widths (restored)
            total_attendance_conducted = 0
            total_attendance_present = 0
            total_marks_sum = 0
            subjects_sorted = sorted(subjects, key=lambda s: 1 if s.get('is_lab', False) else 0)
            # We'll enforce S.No. width after table completion
            for idx, subject in enumerate(subjects_sorted):
                data_row = table.add_row()
                data_cells = data_row.cells
                attendance_conducted = subject['attendance_conducted']
                attendance_present = subject['attendance_present']
                total_attendance_conducted += attendance_conducted
                total_attendance_present += attendance_present
                is_lab = bool(subject.get('is_lab', False))
                if is_lab:
                    dt_marks = '-'
                    st_marks = '-'
                    at_marks = '-'
                    total_marks = '-'
                else:
                    dt_marks = subject['dt_marks'] * 20 / 30
                    st_marks = subject['st_marks']
                    at_marks = subject['at_marks']
                    total_marks = dt_marks + st_marks + at_marks
                    total_marks_sum += total_marks
                row_data = [
                    str(idx + 1),
                    subject['subject_name'],
                    str(attendance_conducted),
                    str(attendance_present),
                    (str(round(dt_marks)) if dt_marks != '-' else '-'),
                    (str(st_marks) if st_marks != '-' else '-'),
                    (str(at_marks) if at_marks != '-' else '-'),
                    (str(round(total_marks)) if total_marks != '-' else '-')
                ]
                for i, data in enumerate(row_data):
                    if i < len(data_cells):
                        data_cells[i].text = data
                        for paragraph in data_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(9)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if is_lab:
                    try:
                        merged = data_cells[4].merge(data_cells[5])
                        merged = merged.merge(data_cells[6])
                        merged = merged.merge(data_cells[7])
                        merged.text = '-'
                        for paragraph in merged.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(9)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception:
                        pass
            total_row = table.add_row()
            total_cells = total_row.cells
            total_cells[1].text = "TOTAL"
            total_cells[2].text = str(total_attendance_conducted)
            total_cells[3].text = str(total_attendance_present)
            try:
                merged_total = total_cells[4].merge(total_cells[5])
                merged_total = merged_total.merge(total_cells[6])
                merged_total = merged_total.merge(total_cells[7])
                merged_total.text = ''
            except Exception:
                pass
            overall_attendance_percent = (total_attendance_present / total_attendance_conducted * 100) if total_attendance_conducted > 0 else 0
            percent_row = table.add_row()
            percent_cells = percent_row.cells
            percent_cells[1].text = "Percentage"
            # Merge attendance columns for percentage value
            percent_row_idx = len(table.rows) - 1
            merged_attendance_cell = table.cell(percent_row_idx, 2)
            merged_attendance_cell.merge(table.cell(percent_row_idx, 3))
            merged_attendance_cell.text = f"{overall_attendance_percent:.2f}%"
            try:
                merged_marks_cell = table.cell(percent_row_idx, 4).merge(table.cell(percent_row_idx, 5))
                merged_marks_cell = merged_marks_cell.merge(table.cell(percent_row_idx, 6))
                merged_marks_cell = merged_marks_cell.merge(table.cell(percent_row_idx, 7))
                merged_marks_cell.text = '-'
                for p in merged_marks_cell.paragraphs:
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(9)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
            for row in [total_row, percent_row]:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(9)
                            run.font.bold = True
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Enforce compact widths at the end to match reference image
            try:
                table.autofit = False
                # Set specific column widths as per specifications
                table.columns[0].width = Inches(0.1)  # S.No. = 0.1 cm
                table.columns[1].width = Inches(0.3)  # Course Title = 3 times S.No. (0.3 cm)
                table.columns[2].width = Inches(0.5)  # Attendance = 5 times S.No. (0.5 cm)
                table.columns[3].width = Inches(0.5)  # Attendance = 5 times S.No. (0.5 cm)
                # CIE-1 Marks columns = 5 times S.No. (0.5 cm)
                table.columns[4].width = Inches(0.5)  # DT (20)
                table.columns[5].width = Inches(0.5)  # ST (10)
                table.columns[6].width = Inches(0.5)  # AT (10)
                table.columns[7].width = Inches(0.5)  # Total (40)
            except Exception:
                # Fallback for individual cell width setting
                for r in table.rows:
                    r.cells[0].width = Inches(0.1)  # S.No. = 0.1 cm
                    r.cells[1].width = Inches(0.3)  # Course Title = 3 times S.No. (0.3 cm)
                    r.cells[2].width = Inches(0.5)  # Attendance = 5 times S.No. (0.5 cm)
                    r.cells[3].width = Inches(0.5)  # Attendance = 5 times S.No. (0.5 cm)
                    r.cells[4].width = Inches(0.5)  # CIE-1 Marks = 5 times S.No. (0.5 cm)
                    r.cells[5].width = Inches(0.5)  # CIE-1 Marks = 5 times S.No. (0.5 cm)
                    r.cells[6].width = Inches(0.5)  # CIE-1 Marks = 5 times S.No. (0.5 cm)
                    r.cells[7].width = Inches(0.5)  # CIE-1 Marks = 5 times S.No. (0.5 cm)
            if template == "Detailed":
                legend_para = doc.add_paragraph()
                legend_text = "*DT – Descriptive Test  ST-Surprise Test  AT- Assignment"
                legend_run = legend_para.add_run(legend_text)
                legend_run.font.name = 'Times New Roman'
                legend_run.font.size = Pt(9)
                attendance_note = doc.add_paragraph()
                attendance_status = "Poor" if overall_attendance_percent < 75 else "Satisfactory"
                note_text = f"Your ward's attendance is {overall_attendance_percent:.2f}% which is {attendance_status}."
                note_run = attendance_note.add_run(note_text)
                note_run.font.name = 'Times New Roman'
                note_run.font.size = Pt(9)
                note_run.font.bold = True
                if overall_attendance_percent < 75:
                    note_run.font.color.rgb = RGBColor(255, 0, 0)
                if include_notes:
                    important_para = doc.add_paragraph()
                    important_run = important_para.add_run("Important Note:")
                    important_run.font.name = 'Times New Roman'
                    important_run.font.size = Pt(10)
                    important_run.font.bold = True
                    notes_text = "→ As per the Osmania University rules, a student must have minimum attendance of 75% in aggregate of all the subjects to be eligible or promoted for the next year. Students having less than 75% attendance in aggregate will not be issued Hall Ticket for the examination, such students will come under Condonation/Detention category.\n"
                    notes_text += "→ As per State Government rules, the student is not eligible for Scholarship if the attendance is less than 75%."
                    notes_para = doc.add_paragraph()
                    notes_run = notes_para.add_run(notes_text)
                    notes_run.font.name = 'Times New Roman'
                    notes_run.font.size = Pt(9)
                if include_backlog:
                    backlog_para = doc.add_paragraph()
                    backlog_run = backlog_para.add_run("Backlog Data:")
                    backlog_run.font.name = 'Times New Roman'
                    backlog_run.font.size = Pt(10)
                    backlog_run.font.bold = True
                    backlog_table = doc.add_table(rows=2, cols=4)
                    backlog_table.style = 'Table Grid'
                    backlog_headers = ['I Sem.', 'II Sem.', 'III Sem.', 'Remarks by Head of the Department']
                    backlog_hdr_cells = backlog_table.rows[0].cells
                    for i, header in enumerate(backlog_headers):
                        backlog_hdr_cells[i].text = header
                        for paragraph in backlog_hdr_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.bold = True
                                run.font.size = Pt(9)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    backlog_data_cells = backlog_table.rows[1].cells
                    for cell in backlog_data_cells:
                        cell.text = "-"
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(9)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
                signature_para = doc.add_paragraph()
                signature_text = "Sign. of the student: ________________    Sign. of the Parent/Guardian: ________________"
                signature_run = signature_para.add_run(signature_text)
                signature_run.font.name = 'Times New Roman'
                signature_run.font.size = Pt(9)
                # doc.add_paragraph()
                final_signature_para = doc.add_paragraph()
                final_signature_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                mentor_run = final_signature_para.add_run("Mentor")
                mentor_run.font.name = 'Times New Roman'
                mentor_run.font.size = Pt(9)
                spaces = " " * 60
                final_signature_para.add_run(spaces)
                hod_run = final_signature_para.add_run("Head of the Department")
                hod_run.font.name = 'Times New Roman'
                hod_run.font.size = Pt(9)
    return doc

def generate_comprehensive_reports(all_students, subjects_data, department_name, report_date, academic_year, semester, attendance_start="", attendance_end="", template="Detailed", include_backlog=True, include_notes=True):
    """Generate comprehensive reports for all students in parallel"""
    individual_reports = {}
    with concurrent.futures.ThreadPoolExecutor() as executor:
        future_to_roll = {
            executor.submit(generate_student_reports, roll, subjects_data, department_name, report_date, academic_year, semester, attendance_start, attendance_end, template, include_backlog, include_notes): roll
            for roll in all_students
        }
        for future in concurrent.futures.as_completed(future_to_roll):
            roll = future_to_roll[future]
            try:
                report = future.result()
                if report:
                    student_name = list(report.keys())[0].split('_Comprehensive_Report_docx')[0]
                    individual_reports[roll] = {
                        'docx_content': report[f"{student_name}_Comprehensive_Report_docx"],
                        'student_name': student_name
                    }
            except Exception as e:
                st.error(f"Error generating report for {roll}: {str(e)}")
    consolidated_doc = create_consolidated_all_students_report(all_students, subjects_data, department_name, report_date, academic_year, semester, attendance_start, attendance_end, template, include_backlog, include_notes)
    consolidated_buffer = BytesIO()
    consolidated_doc.save(consolidated_buffer)
    consolidated_buffer.seek(0)
    return individual_reports, consolidated_buffer.getvalue()
